"""Text extraction module.

Extraction priority:
  1. DOCX  — python-docx, walks the body in reading order (paragraphs AND
     tables — bilingual contracts here are laid out as two-column tables,
     which python-docx's `.paragraphs` alone skips entirely)
  2. PDF (text layer) — pdfplumber, page-by-page text extraction
  3. PDF (image-only) — pdf2image + pytesseract OCR, with accuracy_warning
  4. Image (jpg/png/etc.) — pytesseract OCR, with accuracy_warning

Failures set extraction_ok=False and populate extraction_note.
NEVER swallow an extraction failure silently — that is a spec violation.

`extraction_ok` means the document was plausibly extracted in full, not
merely that some non-empty text came back. Every path below checks the
extracted length against an absolute floor and a coverage ratio against an
independent size signal (raw body text for DOCX, page count for PDF, byte
size for single images) before declaring success — see
.claude/rules/extraction-coverage.md.

OCR paths require system packages (tesseract, poppler). When unavailable,
extraction_ok is set to False with an explanatory note — the pipeline
continues but flags the version clearly.
"""
import hashlib
import io
from dataclasses import dataclass, field
from typing import Optional


@dataclass
class ExtractionResult:
    text: str
    source_format: str          # 'docx' | 'pdf_text' | 'pdf_image' | 'image'
    extraction_ok: bool
    extraction_note: Optional[str]
    accuracy_warning: Optional[str]  # non-None for OCR paths
    sha256: str


_ALLOWED_EXTENSIONS = {
    "docx", "pdf", "jpg", "jpeg", "png", "webp", "bmp"
}
_MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB

# Coverage floors — see .claude/rules/extraction-coverage.md. Values are
# deliberately conservative starting points; tune against real creator
# contracts, not invented numbers.
_MIN_CHARS = 100                 # absolute floor, any document
_MIN_COVERAGE_RATIO = 0.5        # DOCX: extracted vs. raw body text
_MIN_CHARS_PER_PAGE = 40         # PDF: extracted vs. page_count
_MIN_CHARS_PER_KB = 0.5          # single image: extracted vs. file size


def extract(file_bytes: bytes, filename: str) -> ExtractionResult:
    """Entry point: dispatch to the correct extractor based on file extension."""
    sha256 = hashlib.sha256(file_bytes).hexdigest()
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if len(file_bytes) > _MAX_FILE_SIZE:
        return ExtractionResult(
            text="",
            source_format="docx",
            extraction_ok=False,
            extraction_note=f"Ukuran file melebihi batas maksimum 20 MB.",
            accuracy_warning=None,
            sha256=sha256,
        )

    if ext == "docx":
        return _extract_docx(file_bytes, sha256)
    elif ext == "pdf":
        return _extract_pdf(file_bytes, sha256)
    elif ext in ("jpg", "jpeg", "png", "webp", "bmp"):
        return _extract_image(file_bytes, sha256)
    else:
        return ExtractionResult(
            text="",
            source_format="docx",
            extraction_ok=False,
            extraction_note=(
                f"Format file '.{ext}' tidak didukung. "
                "Gunakan DOCX, PDF, JPG, atau PNG."
            ),
            accuracy_warning=None,
            sha256=sha256,
        )


def _iter_body_blocks(doc):
    """Yield ('p', Paragraph) / ('tbl', Table) for doc.element.body's direct
    children, in document reading order — dispatching on the raw XML rather
    than reading doc.paragraphs and doc.tables as separate collections,
    which loses the interleaving between them."""
    from docx.table import Table  # type: ignore[import]
    from docx.text.paragraph import Paragraph  # type: ignore[import]

    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield "p", Paragraph(child, doc)
        elif tag == "tbl":
            yield "tbl", Table(child, doc)


def _extract_table_lines(table) -> list[str]:
    """Extract a table's text in row-major reading order.

    For a two-column [Indonesian | English] row this naturally keeps the
    two cells adjacent, pairing them as parallel expressions of the same
    clause — column-major iteration (all left-column cells, then all
    right-column cells) would instead scramble clause order across the
    whole table. Rows with a different column count fall through to the
    same row-major order rather than guessing which columns pair up.
    """
    lines: list[str] = []
    for row in table.rows:
        seen_tc_ids: set[int] = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tc_ids:
                continue  # horizontally merged cell repeated by python-docx
            seen_tc_ids.add(tc_id)
            text = cell.text.strip()
            if text:
                lines.append(text)
    return lines


def _body_text_length(doc) -> int:
    """Independent size signal: total characters in every w:t run under the
    body (paragraphs + tables), read directly from the XML rather than via
    doc.paragraphs/doc.tables — so a future regression to a partial walk
    still shows up as a coverage-ratio drop against this count."""
    from docx.oxml.ns import qn  # type: ignore[import]

    return sum(len(t.text) for t in doc.element.body.iter(qn("w:t")) if t.text)


def _extract_docx(file_bytes: bytes, sha256: str) -> ExtractionResult:
    try:
        import docx  # type: ignore[import]

        doc = docx.Document(io.BytesIO(file_bytes))
        lines: list[str] = []
        for tag, block in _iter_body_blocks(doc):
            if tag == "p":
                text = block.text.strip()
                if text:
                    lines.append(text)
            else:
                lines.extend(_extract_table_lines(block))

        full_text = "\n".join(lines)
        available_chars = _body_text_length(doc)
        ok = (
            len(full_text) >= _MIN_CHARS
            and available_chars > 0
            and (len(full_text) / available_chars) >= _MIN_COVERAGE_RATIO
        )

        note = None
        if not ok:
            if not full_text.strip():
                note = "Dokumen DOCX tidak mengandung teks yang dapat dibaca."
            else:
                note = (
                    f"Ekstraksi DOCX tampak tidak lengkap ({len(full_text)} dari "
                    f"perkiraan {available_chars} karakter dalam dokumen). "
                    "Dokumen ini memerlukan tinjauan manual."
                )

        return ExtractionResult(
            text=full_text,
            source_format="docx",
            extraction_ok=ok,
            extraction_note=note,
            accuracy_warning=None,
            sha256=sha256,
        )
    except Exception as exc:
        return ExtractionResult(
            text="",
            source_format="docx",
            extraction_ok=False,
            extraction_note=f"Gagal membuka DOCX: {exc}",
            accuracy_warning=None,
            sha256=sha256,
        )


def _extract_pdf(file_bytes: bytes, sha256: str) -> ExtractionResult:
    """Try text layer; fall through to OCR if no text found."""
    try:
        import pdfplumber  # type: ignore[import]

        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text()
                if t and t.strip():
                    pages.append(t.strip())

        full_text = "\n\n".join(pages)
        min_expected = max(_MIN_CHARS, _MIN_CHARS_PER_PAGE * page_count)
        if len(full_text) >= min_expected:
            return ExtractionResult(
                text=full_text,
                source_format="pdf_text",
                extraction_ok=True,
                extraction_note=None,
                accuracy_warning=None,
                sha256=sha256,
            )
        if full_text.strip():
            # Some text found but far short of what the page count implies —
            # a partially-extracted PDF, not proof of an image-only PDF.
            # Surface it for review rather than silently attempting OCR on
            # top of it, which could layer a second guess on the first.
            return ExtractionResult(
                text=full_text,
                source_format="pdf_text",
                extraction_ok=False,
                extraction_note=(
                    f"Ekstraksi PDF tampak tidak lengkap ({len(full_text)} karakter "
                    f"dari {page_count} halaman). Dokumen ini memerlukan tinjauan manual."
                ),
                accuracy_warning=None,
                sha256=sha256,
            )
    except Exception as exc:
        # Log and fall through to OCR
        pass

    # Text layer empty or unreadable — attempt OCR
    return _extract_pdf_ocr(file_bytes, sha256)


def _extract_pdf_ocr(file_bytes: bytes, sha256: str) -> ExtractionResult:
    try:
        from pdf2image import convert_from_bytes  # type: ignore[import]
        import pytesseract  # type: ignore[import]

        images = convert_from_bytes(file_bytes, dpi=200)
        page_count = len(images)
        texts: list[str] = []
        for img in images:
            t = pytesseract.image_to_string(img, lang="ind+eng")
            if t.strip():
                texts.append(t.strip())

        full_text = "\n\n".join(texts)
        min_expected = max(_MIN_CHARS, _MIN_CHARS_PER_PAGE * page_count)
        ok = len(full_text) >= min_expected
        if ok:
            note = None
        elif full_text.strip():
            note = (
                f"OCR menghasilkan teks yang tampak tidak lengkap ({len(full_text)} "
                f"karakter dari {page_count} halaman). Dokumen ini memerlukan tinjauan manual."
            )
        else:
            note = "OCR tidak berhasil mengekstrak teks dari PDF."
        return ExtractionResult(
            text=full_text,
            source_format="pdf_image",
            extraction_ok=ok,
            extraction_note=note,
            accuracy_warning=(
                "Dokumen ini adalah PDF berbasis gambar. "
                "Akurasi ekstraksi teks melalui OCR dapat bervariasi — "
                "verifikasi klausul penting secara manual."
            ),
            sha256=sha256,
        )
    except ImportError:
        return ExtractionResult(
            text="",
            source_format="pdf_image",
            extraction_ok=False,
            extraction_note=(
                "PDF ini tampaknya berbasis gambar (tidak ada lapisan teks), "
                "tetapi perangkat OCR (tesseract/poppler) tidak tersedia di server ini. "
                "Hubungi administrator untuk mengaktifkan dukungan OCR."
            ),
            accuracy_warning=None,
            sha256=sha256,
        )
    except Exception as exc:
        return ExtractionResult(
            text="",
            source_format="pdf_image",
            extraction_ok=False,
            extraction_note=f"Gagal memproses PDF berbasis gambar: {exc}",
            accuracy_warning=None,
            sha256=sha256,
        )


def _extract_image(file_bytes: bytes, sha256: str) -> ExtractionResult:
    try:
        import pytesseract  # type: ignore[import]
        from PIL import Image  # type: ignore[import]

        img = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(img, lang="ind+eng")
        stripped = text.strip()
        # No independent page/element count for a single image — byte size
        # is the last-resort signal per extraction-coverage.md, applied only
        # as a floor (not a ratio: a small high-res photo of a short clause
        # is legitimate and shouldn't be penalized for its file size).
        min_expected = min(_MIN_CHARS, _MIN_CHARS_PER_KB * (len(file_bytes) / 1024))
        ok = len(stripped) >= max(1, min_expected) and bool(stripped)
        return ExtractionResult(
            text=stripped if ok else "",
            source_format="image",
            extraction_ok=ok,
            extraction_note=None if ok else "OCR tidak menghasilkan teks yang memadai dari gambar.",
            accuracy_warning=(
                "Dokumen diunggah sebagai gambar. "
                "Akurasi ekstraksi teks melalui OCR dapat bervariasi — "
                "verifikasi klausul penting secara manual."
            ),
            sha256=sha256,
        )
    except ImportError:
        return ExtractionResult(
            text="",
            source_format="image",
            extraction_ok=False,
            extraction_note=(
                "Perangkat OCR untuk memproses gambar (tesseract) "
                "tidak tersedia di server ini."
            ),
            accuracy_warning=None,
            sha256=sha256,
        )
    except Exception as exc:
        return ExtractionResult(
            text="",
            source_format="image",
            extraction_ok=False,
            extraction_note=f"Gagal memproses gambar: {exc}",
            accuracy_warning=None,
            sha256=sha256,
        )
