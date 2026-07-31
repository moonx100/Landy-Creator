"""Tests for landy.extraction — table-aware DOCX walk + coverage floor.

Regression coverage for the bug where `_extract_docx` read only
`doc.paragraphs`, skipping the `w:tbl` elements that Indonesian bilingual
contracts use to lay out their entire body. On the fixture this produced
51 characters of extracted text out of ~60,000 actually present.
"""
from __future__ import annotations

import re
from pathlib import Path

import pytest

from landy.extraction import (
    _MIN_CHARS,
    _MIN_COVERAGE_RATIO,
    _body_text_length,
    extract,
)

FIXTURE_PATH = (
    Path(__file__).resolve().parents[3]
    / "tests"
    / "fixtures"
    / "SAMPLE_-_Sale_Purchase_Software_-_OriClient.docx"
)


class TestDocxTableExtraction:
    def test_full_bilingual_contract_is_extracted(self):
        assert FIXTURE_PATH.exists(), f"fixture missing at {FIXTURE_PATH}"
        data = FIXTURE_PATH.read_bytes()

        result = extract(data, FIXTURE_PATH.name)

        assert result.extraction_ok is True
        assert result.extraction_note is None
        # Historical bug extracted 51 chars; full contract is ~60k.
        assert len(result.text) > 10_000

        pasal_numbers = sorted(
            {int(m) for m in re.findall(r"PASAL (\d+)", result.text)}
        )
        assert pasal_numbers == list(range(1, 19))

        assert "kerahasiaan" in result.text.lower()
        assert "bahasa" in result.text.lower()

    def test_synthetic_near_empty_extraction_of_large_document_is_caught(self):
        """Reproduces the exact regression class the coverage floor guards
        against: reading only paragraphs on a table-based document yields a
        near-empty string that must NOT pass the coverage check, even though
        it's non-empty."""
        import docx

        assert FIXTURE_PATH.exists(), f"fixture missing at {FIXTURE_PATH}"
        doc = docx.Document(str(FIXTURE_PATH))

        paragraph_only_text = "\n".join(
            p.text.strip() for p in doc.paragraphs if p.text.strip()
        )
        available_chars = _body_text_length(doc)

        # The historical bug: ~51 chars extracted from a document with tens
        # of thousands of characters actually present.
        assert len(paragraph_only_text) < 200
        assert available_chars > 10_000

        coverage_ok = (
            len(paragraph_only_text) >= _MIN_CHARS
            and available_chars > 0
            and (len(paragraph_only_text) / available_chars) >= _MIN_COVERAGE_RATIO
        )
        assert coverage_ok is False
